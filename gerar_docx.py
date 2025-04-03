from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd

def gerar_documento(df: pd.DataFrame, nome_arquivo: str):
    doc = Document()

    # Estilo do título
    titulo = doc.add_heading("Registro de Aquisição - Biblioteca", 0)
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    for index, row in df.iterrows():
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        p.add_run("Título: ").bold = True
        p.add_run(f"{row['Título']}\n")

        p.add_run("Autor: ").bold = True
        p.add_run(f"{row['Autor']}\n")

        p.add_run("ISBN: ").bold = True
        p.add_run(f"{row['ISBN']}\n")

        p.add_run("ID do Acervo: ").bold = True
        p.add_run(f"{row['ID_Acervo']}\n")

        p.add_run("Exemplar: ").bold = True
        p.add_run(f"{row['Exemplar']}\n")

        p.add_run("Código Cutter: ").bold = True
        p.add_run(f"{row['Código Cutter']}\n")

        p.add_run("Ano de Publicação: ").bold = True
        p.add_run(f"{row['Ano de Publicação']}\n")

        p.add_run("Classificação CDD: ").bold = True
        p.add_run(f"{row['Classificação CDD']}\n")

        doc.add_paragraph("\n")

    doc.save(nome_arquivo)
